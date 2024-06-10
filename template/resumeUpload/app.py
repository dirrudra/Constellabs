from flask import Flask, render_template, request, redirect, url_for
import os
from werkzeug.utils import secure_filename
from docx import Document
from PyPDF2 import PdfReader
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
SAVED_TEXTS_FOLDER = 'saved_texts'
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['SAVED_TEXTS_FOLDER'] = SAVED_TEXTS_FOLDER

# Configure Google Generative AI

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# genai.configure(api_key="AIzaSyDq24IX22APnfwraBYjvxK5znvqTsT1L2o")

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_docx(filepath):
    try:
        doc = Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from {filepath}: {e}")
        return ""

def extract_text_from_pdf(filepath):
    try:
        text = ""
        with open(filepath, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                else:
                    print(f"Warning: Page {reader.pages.index(page)} has no extractable text.")
        return text
    except Exception as e:
        print(f"Error extracting text from {filepath}: {e}")
        return ""

def generate_content_from_text(input_text):
    chat_session = model.start_chat(history=[])
    print(input_text,"session logs")
    response = chat_session.send_message("Generate a simple website with proper formatting \n Do not give any descriptions, generate only HTML CSS with clean animations for this using and also NO DESCRIPTIONS OR EXPLNATIONS !!:\n\n\n"+input_text)
    return response.text

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            if filename.rsplit('.', 1)[1].lower() == 'docx':
                text = extract_text_from_docx(filepath)
            elif filename.rsplit('.', 1)[1].lower() == 'pdf':
                text = extract_text_from_pdf(filepath)
            else:
                print("Unsupported file format.")
                return redirect(request.url)

            if text:
                generated_content = generate_content_from_text(text)
                text_filename = filename.rsplit('.', 1)[0] + '_generated.txt'
                with open(os.path.join(app.config['SAVED_TEXTS_FOLDER'], text_filename), 'w') as text_file:
                    text_file.write(generated_content)
                return render_template('index.html', generated_content=generated_content)
            else:
                print(f"Failed to extract text from {filepath}")

            return redirect(url_for('upload_file'))

    return render_template('index.html')

@app.route('/generated_website', methods=['GET'])
def generated_website():
    generated_content = request.args.get('generated_content', '')
    print(generated_content)
    # return res.jsonify({'generated_content': generated_content})
    return render_template('generated_website.html', generated_content=generated_content)

if __name__ == '__main__':
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(SAVED_TEXTS_FOLDER, exist_ok=True)
    app.run(debug=True)


# from flask import Flask, render_template, request, redirect, url_for
# import os
# from werkzeug.utils import secure_filename
# from docx import Document

# app = Flask(__name__)

# UPLOAD_FOLDER = 'uploads'
# SAVED_TEXTS_FOLDER = 'saved_texts'
# ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}

# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# app.config['SAVED_TEXTS_FOLDER'] = SAVED_TEXTS_FOLDER

# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_docx(filepath):
#     try:
#         doc = Document(filepath)
#         full_text = []
#         for para in doc.paragraphs:
#             full_text.append(para.text)
#         return '\n'.join(full_text)
#     except Exception as e:
#         print(f"Error extracting text from {filepath}: {e}")
#         return ""

# @app.route('/', methods=['GET', 'POST'])
# def upload_file():
#     if request.method == 'POST':
#         if 'file' not in request.files:

#             return redirect(request.url)
#         file = request.files['file']
#         if file.filename == '':

#             return redirect(request.url)
#         if file and allowed_file(file.filename):
#             print("came here...")
#             filename = secure_filename(file.filename)
#             filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#             file.save(filepath)

#             if filename.rsplit('.', 1)[1].lower() == 'docx':
#                 text = extract_text_from_docx(filepath)
#                 if text:
#                     text_filename = filename.rsplit('.', 1)[0] + '.txt'
#                     with open(os.path.join(app.config['SAVED_TEXTS_FOLDER'], text_filename), 'w') as text_file:
#                         print(text)
#                         text_file.write(text)
#                 else:
#                     print(f"Failed to extract text from {filepath}")

#             return redirect(url_for('upload_file'))

#     return render_template('index.html')

# if __name__ == '__main__':
#     os.makedirs(UPLOAD_FOLDER, exist_ok=True)
#     os.makedirs(SAVED_TEXTS_FOLDER, exist_ok=True)
#     app.run(debug=True)


